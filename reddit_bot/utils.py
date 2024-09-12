import os
import getpass
from langchain.chains.llm import LLMChain
from langchain.chains.summarize.refine_prompts import prompt_template
from langchain_openai import ChatOpenAI
from qdrant_client import QdrantClient
import bs4
from langchain import hub
from langchain_community.document_loaders import WebBaseLoader
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import RunnablePassthrough
from langchain_openai import OpenAIEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Qdrant
import praw
from datetime import datetime
from django.http import JsonResponse
from openai import OpenAI
from reddit_bot import keys
from langchain.prompts import PromptTemplate
# from .models import RedditPost, GeneratedResponse
from docx import Document as DocxDocument
from typing import Dict


QDRANT_API_KEY = keys.QDRANT_API_KEY
QDRANT_URL = keys.QDRANT_URL


def create_collection(splited_text):
    """
    Create a Qdrant collection from a list of text chunks.
    """
    docs = splited_text
    embeddings=OpenAIEmbeddings()

    qdrant = Qdrant.from_documents(
        docs,
        embeddings,
        url=QDRANT_URL,
        api_key=QDRANT_API_KEY,
        collection_name="hongyuan",
    )

    return qdrant


class Document:
    def __init__(self, page_content, metadata):
        self.page_content = page_content
        self.metadata = metadata


def read_word_document(file_path):
    doc = DocxDocument(file_path)

    page_content = '\n'.join([para.text for para in doc.paragraphs if para.text])

    metadata = {
        'title': doc.core_properties.title,
        'author': doc.core_properties.author,
        'created': doc.core_properties.created,
        'subject': doc.core_properties.subject
    }

    return Document(page_content, metadata)

def parse_retriever_input(params: Dict):
    return params["messages"][-1].content

